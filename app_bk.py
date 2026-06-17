import streamlit as st
import pandas as pd
import json
import os
import datetime
import random
import calendar
import io
import re
from config import Config
import db

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Khởi tạo database và bảng
db.init_db()

# Thiết lập cấu hình trang
st.set_page_config(
    page_title=Config.APP_NAME,
    page_icon="🏠",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS để giao diện chuyên nghiệp hơn theo brand leducnam.com
# Danh sách trạng thái chuẩn
STATUS_LIST = ["Mới tạo", "Đã tiếp nhận", "Đang xử lý", "Chờ xử lý", "Hoàn thành", "Từ chối"]
STATUS_COLORS = {
    "Mới tạo": "#3B82F6",       # Xanh dương
    "Đã tiếp nhận": "#8B5CF6",  # Tím
    "Đang xử lý": "#F59E0B",    # Vàng cam
    "Chờ xử lý": "#EF4444",     # Đỏ
    "Hoàn thành": "#10B981",     # Xanh lá
    "Từ chối": "#6B7280",       # Xám
}

def status_badge(status_text):
    color = STATUS_COLORS.get(status_text, "#6B7280")
    return f'<span style="background:{color};color:#fff;padding:2px 10px;border-radius:12px;font-size:0.85em;font-weight:600;">{status_text}</span>'

st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1E3A8A;
        font-weight: 700;
        margin-bottom: 0.5rem;
    }
    .sub-header {
        font-size: 1.2rem;
        color: #4B5563;
        margin-bottom: 2rem;
    }
    .stAppHeader .stToolbarActions {display:none;}
</style>
""", unsafe_allow_html=True)

# Lấy dữ liệu từ database
db_tickets = db.get_all_tickets()
db_tasks = db.get_tasks()

# -----------------
# XỬ LÝ DEEP LINKING (URL Query Parameters)
# -----------------
url_ticket_id = st.query_params.get("ticket")
if url_ticket_id and st.session_state.get("view_mode") != "full":
    st.session_state["view_mode"] = "only"
    st.session_state["active_ticket_id"] = url_ticket_id

# Hàm sinh mã Ticket tự động
def gen_ticket_id():
    now = datetime.datetime.now()
    return f"GREE-IT_{now.strftime('%Y%m')}_{random.randint(100, 999)}"

FORM_CONFIGS = [
    {
        "label": "Form 1: Khai báo Model phân loại chi phí bảo hành",
        "type": "Khai_Bao_Model_Bao_Hanh",
        "title": "Danh sách Model đã cập nhật hệ thống",
        "report_fields": ["model_name"],
        "columns": [
            ("Mã Ticket", "id"),
            ("Tên model", "model_name"),
            ("Loại chi phí", "cost_type"),
            ("Công suất", "capacity_range"),
            ("Loại sản phẩm", "product_type"),
            ("Ngày y/c", "created_at"),
            ("Người y/c", "requester"),
            ("Trạng thái", "status"),
            ("Ghi chú", "note"),
        ],
    },
    {
        "label": "Form 2: Khai báo mã linh kiện mới",
        "type": "Khai_Bao_Ma_Linh_Kien",
        "title": "Danh sách mã linh kiện đã cập nhật hệ thống",
        "report_fields": ["part_code", "part_name_vi"],
        "columns": [
            ("Mã Ticket", "id"),
            ("Mã linh kiện", "part_code"),
            ("Tên linh kiện (EN)", "part_name_en"),
            ("Tên linh kiện (VI)", "part_name_vi"),
            ("Ngày y/c", "created_at"),
            ("Người y/c", "requester"),
            ("Trạng thái", "status"),
            ("Mô tả", "description"),
        ],
    },
    {
        "label": "Form 3: Yêu cầu điều chỉnh tồn kho hệ thống",
        "type": "Yeu_Cau_Dieu_Chinh_Ton_Kho",
        "title": "Danh sách yêu cầu điều chỉnh tồn kho",
        "report_fields": ["note", "warehouse"],
        "columns": [
            ("Mã Ticket", "id"),
            ("Kho/Trạm", "warehouse"),
            ("Mã linh kiện", "part_code"),
            ("Tính chất LK", "part_nature"),
            ("Phiếu xuất", "export_voucher"),
            ("SL điều chỉnh", "adjusted_quantity"),
            ("Ngày y/c", "created_at"),
            ("Người y/c", "requester"),
            ("Trạng thái", "status"),
            ("Ghi chú", "note"),
        ],
    },
    {
        "label": "Form 4: Đăng ký thông tin trạm bảo hành mới",
        "type": "Dang_Ky_Tram_Bao_Hanh_Moi",
        "title": "Danh sách trạm bảo hành đăng ký mới",
        "report_fields": ["company_info.name"],
        "columns": [
            ("Mã Ticket", "id"),
            ("Tên công ty", "company_info.name"),
            ("Mã số thuế", "company_info.tax_code"),
            ("Email", "company_info.email"),
            ("Điện thoại", "company_info.phone"),
            ("User hệ thống", "company_info.system_username"),
            ("Số KTV", "technicians"),
            ("Số kho", "associated_warehouses"),
            ("Ngày y/c", "created_at"),
            ("Người y/c", "requester"),
            ("Trạng thái", "status"),
        ],
    },
    {
        "label": "Form 5: Đăng ký tài khoản user nội bộ Gree",
        "type": "Dang_Ky_Tai_Khoan_User_Noi_Bo",
        "title": "Danh sách tài khoản user nội bộ đăng ký mới",
        "report_fields": ["full_name", "user_group"],
        "columns": [
            ("Mã Ticket", "id"),
            ("Họ tên", "full_name"),
            ("Điện thoại", "phone"),
            ("Email Gree", "company_email"),
            ("Nhóm user", "user_group"),
            ("Line Call Center", "call_center_line"),
            ("Link chính", "main_link"),
            ("Username test", "test_account.username"),
            ("Password test", "test_account.password"),
            ("Link test", "test_account.test_link"),
            ("Ngày y/c", "created_at"),
            ("Người y/c", "requester"),
            ("Trạng thái", "status"),
        ],
    },
    {
        "label": "Form 6: Khai báo Model cho Import Hồ sơ máy",
        "type": "Khai_Bao_Model_Ho_So_May",
        "title": "Danh sách Model hồ sơ máy đã cập nhật hệ thống",
        "report_fields": ["model_name"],
        "columns": [
            ("Mã Ticket", "id"),
            ("Tên model", "model_name"),
            ("Loại máy", "machine_type"),
            ("BH Máy (tháng)", "warranty_months_machine"),
            ("BH Block (tháng)", "warranty_months_compressor"),
            ("Ngày y/c", "created_at"),
            ("Người y/c", "requester"),
            ("Trạng thái", "status"),
        ],
    },
    {
        "label": "Form 7: Admin ghi nhận nội dung hỗ trợ (Admin Logs)",
        "type": "Admin_Web_Noted_Log_Ho_Tro",
        "title": "Danh sách log hỗ trợ Admin Web",
        "report_fields": ["request_detail"],
        "columns": [
            ("Mã Ticket", "id"),
            ("Mã ca", "case_code"),
            ("Người xử lý IT", "it_assignee"),
            ("Ngày hoàn thành", "completion_date"),
            ("Ngày y/c", "created_at"),
            ("Người y/c", "requester"),
            ("Trạng thái", "status"),
            ("Nội dung yêu cầu", "request_detail"),
            ("Log xử lý nội bộ", "internal_action_log"),
        ],
    },
    {
        "label": "Form 8: Import bảng giá linh kiện",
        "type": "Import_Bang_Gia_Linh_Kien",
        "title": "Danh sách giá linh kiện đã cập nhật hệ thống",
        "report_fields": ["part_code", "part_name_vi"],
        "columns": [
            ("Mã Ticket", "id"),
            ("Mã linh kiện", "part_code"),
            ("Tên linh kiện (VI)", "part_name_vi"),
            ("Nhóm Model", "model_group"),
            ("Loại sản phẩm", "product_type"),
            ("Giá có VAT", "price_vat"),
            ("Giá chưa VAT", "price_no_vat"),
            ("Ngày y/c", "created_at"),
            ("Người y/c", "requester"),
            ("Trạng thái", "status"),
        ],
    },
    {
        "label": "Form 9: Khai báo danh mục chi phí",
        "type": "Khai_Bao_Danh_Muc_Chi_Phi",
        "title": "Danh sách danh mục chi phí đã cập nhật hệ thống",
        "report_fields": ["content"],
        "columns": [
            ("Mã Ticket", "id"),
            ("Mã chi phí", "cost_code"),
            ("Nội dung", "content"),
            ("Đơn giá", "unit_price"),
            ("Công suất", "capacity"),
            ("Loại chi phí", "cost_type"),
            ("Nhóm sản phẩm", "product_group"),
            ("Ngày y/c", "created_at"),
            ("Người y/c", "requester"),
            ("Trạng thái", "status"),
        ],
    },
    {
        "label": "Form 10: Yêu cầu tổng hợp (Khác)",
        "type": "Yeu_Cau_Tong_Hop_Khac",
        "title": "Danh sách yêu cầu tổng hợp / không thuộc 9 Form trên",
        "report_fields": ["request_detail"],
        "columns": [
            ("Mã Ticket", "id"),
            ("Tiêu đề yêu cầu", "subject"),
            ("Liên quan Form", "related_form"),
            ("Nội dung yêu cầu", "request_detail"),
            ("Ngày y/c", "created_at"),
            ("Người y/c", "requester"),
            ("Trạng thái", "status"),
        ],
    },
]

FORM_LABEL_TO_TYPE = {f["label"]: f["type"] for f in FORM_CONFIGS}
FORM_TYPE_TO_CONFIG = {f["type"]: f for f in FORM_CONFIGS}
LEGACY_FORM_TYPES = {
    f["label"]: f["type"] for f in FORM_CONFIGS
}

def normalize_form_type(form_type):
    return LEGACY_FORM_TYPES.get(form_type, form_type)

def parse_form_data(ticket):
    f_json = ticket.get("form_data", {})
    if isinstance(f_json, str):
        try:
            return json.loads(f_json)
        except Exception:
            return {}
    return f_json or {}

def tickets_for_form(tickets, form_type):
    return [t for t in tickets if normalize_form_type(t.get("form_type")) == form_type]

def format_ticket_date(value):
    if isinstance(value, datetime.datetime):
        return value.strftime("%d/%m/%Y")
    return value or ""

def to_date(value):
    """Chuẩn hóa 1 giá trị (datetime/date/string) về kiểu date để so sánh theo tuần."""
    if isinstance(value, datetime.datetime):
        return value.date()
    if isinstance(value, datetime.date):
        return value
    if isinstance(value, str) and value.strip():
        for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d", "%d/%m/%Y"):
            try:
                return datetime.datetime.strptime(value.strip(), fmt).date()
            except ValueError:
                continue
    return None

def month_first_monday(year, month):
    """Thứ 2 của tuần làm việc chứa ngày 1 của tháng (có thể rơi vào cuối tháng trước)."""
    first_of_month = datetime.date(year, month, 1)
    return first_of_month - datetime.timedelta(days=first_of_month.weekday())

def week_of_month(d):
    """Số thứ tự tuần làm việc (Thứ 2 → Thứ 7) chứa ngày d, đánh số theo tháng của d.
    Tuần 1 = tuần chứa ngày 1 của tháng."""
    monday0 = month_first_monday(d.year, d.month)
    return (d - monday0).days // 7 + 1

def max_weeks_in_month(year, month):
    last_day = calendar.monthrange(year, month)[1]
    return week_of_month(datetime.date(year, month, last_day))

def week_date_range(year, month, week_no):
    """Trả về (Thứ 2, Thứ 7) của tuần làm việc thứ week_no trong tháng được chọn."""
    monday0 = month_first_monday(year, month)
    monday = monday0 + datetime.timedelta(days=7 * (week_no - 1))
    saturday = monday + datetime.timedelta(days=5)
    return monday, saturday

def get_form_value(data, key):
    value = data
    for part in key.split("."):
        if not isinstance(value, dict):
            return ""
        value = value.get(part)
    if isinstance(value, list):
        return len(value)
    return value if value is not None else ""

def parse_pipe_rows(raw_text, fields):
    rows = []
    for line in raw_text.splitlines():
        line = line.strip()
        if not line:
            continue
        parts = [part.strip() for part in line.split("|")]
        row = {field: parts[idx] if idx < len(parts) else "" for idx, field in enumerate(fields)}
        rows.append(row)
    return rows

def missing_form5_required_fields(form_data):
    required_fields = {
        "Họ tên": form_data.get("full_name"),
        "Điện thoại": form_data.get("phone"),
        "Email Gree": form_data.get("company_email"),
        "Nhóm user": form_data.get("user_group"),
        "Link chính": form_data.get("main_link"),
        "Username test": form_data.get("test_account", {}).get("username"),
        "Password test": form_data.get("test_account", {}).get("password"),
        "Link test": form_data.get("test_account", {}).get("test_link"),
    }
    return [label for label, value in required_fields.items() if not str(value or "").strip()]

def build_form_rows(tickets, config):
    rows = []
    for t in tickets:
        f_json = parse_form_data(t)
        row = {}
        for label, key in config["columns"]:
            if key == "id":
                row[label] = t.get("id", "")
            elif key == "created_at":
                row[label] = format_ticket_date(t.get("created_at"))
            elif key == "requester":
                row[label] = t.get("requester", "")
            elif key == "status":
                row[label] = t.get("status", "Mới tạo")
            elif key == "subject":
                row[label] = t.get("subject", "")
            else:
                row[label] = get_form_value(f_json, key)
        rows.append(row)
    return rows

def generate_weekly_report_docx(ctx):
    """Tạo file Word (.docx) báo cáo tuần từ dữ liệu đã tổng hợp. Trả về buffer BytesIO."""
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    heading = doc.add_heading(
        f"BÁO CÁO CÔNG VIỆC TUẦN {ctx['week']} - THÁNG {ctx['month']}/{ctx['year']}",
        level=1
    )
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p1 = doc.add_paragraph()
    p1.add_run("Người báo cáo: ").bold = True
    p1.add_run(ctx['reporter'])

    p2 = doc.add_paragraph()
    p2.add_run("Khoảng thời gian: ").bold = True
    p2.add_run(f"{ctx['start_date'].strftime('%d/%m/%Y')} - {ctx['end_date'].strftime('%d/%m/%Y')}")

    doc.add_heading("1. Tổng quan số liệu trong tuần", level=2)
    t1 = doc.add_table(rows=1, cols=2)
    t1.style = 'Light Grid Accent 1'
    t1.rows[0].cells[0].text = "Chỉ số"
    t1.rows[0].cells[1].text = "Số lượng"
    for label, value in ctx['metrics']:
        row = t1.add_row().cells
        row[0].text = label
        row[1].text = str(value)

    if ctx['summary_rows']:
        doc.add_heading("2. Ticket mới theo Form", level=2)
        t2 = doc.add_table(rows=1, cols=2)
        t2.style = 'Light Grid Accent 1'
        t2.rows[0].cells[0].text = "Form"
        t2.rows[0].cells[1].text = "Số lượng"
        for row_data in ctx['summary_rows']:
            r = t2.add_row().cells
            r[0].text = row_data["Form"]
            r[1].text = str(row_data["Số lượng"])

    doc.add_heading("3. Công việc đã hoàn thành trong tuần", level=2)
    if ctx['completed_rows']:
        t3 = doc.add_table(rows=1, cols=4)
        t3.style = 'Light Grid Accent 1'
        for i, label in enumerate(["Mã Ticket", "Loại Form", "Người y/c", "Nội dung"]):
            t3.rows[0].cells[i].text = label
        for row_data in ctx['completed_rows']:
            r = t3.add_row().cells
            r[0].text = row_data['id']
            r[1].text = row_data['form_label']
            r[2].text = row_data['requester']
            r[3].text = row_data['subject']
    else:
        doc.add_paragraph("Không có ticket nào hoàn thành trong tuần này.")

    doc.add_heading("4. Công việc bổ sung (nhập tay)", level=2)
    doc.add_paragraph(ctx['manual_tasks'].strip() if ctx['manual_tasks'] else "(Không có)")

    doc.add_heading("5. Vấn đề gặp phải / Lưu ý", level=2)
    doc.add_paragraph(ctx['issues'].strip() if ctx['issues'] else "(Không có)")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Xuất báo cáo theo ĐÚNG MẪU CÔNG TY (tìm & thay nội dung trong file .docx thật) ---

TEMPLATE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates", "BaoCaoTuan_Template.docx")

def build_ticket_report_line(t, cfg):
    """Sinh 1 dòng mô tả ticket trong báo cáo: 'Mã – Nội dung (Người y/c) – Trạng thái.'"""
    f_json = parse_form_data(t)
    detail_parts = []
    for key in cfg.get("report_fields", []):
        val = get_form_value(f_json, key)
        if val not in (None, "", "None"):
            detail_parts.append(str(val))
    detail = " - ".join(detail_parts) if detail_parts else t.get("subject", "")
    requester = t.get("requester", "")
    status = t.get("status", "")
    return f"{t.get('id', '')} – {detail} ({requester}) – {status}."

def build_warranty_section_lines(tickets):
    """Sinh nội dung cột 'Các công việc thực hiện' của dòng HỆ THỐNG BẢO HÀNH, phân theo từng Form.
    Trả về list (text, bold) — mỗi tuple là 1 đoạn (paragraph) riêng."""
    lines = []
    section_no = 0
    for cfg in FORM_CONFIGS:
        form_tickets = [t for t in tickets if normalize_form_type(t.get('form_type')) == cfg['type']]
        if not form_tickets:
            continue
        section_no += 1
        short_label = cfg['label'].split(": ", 1)[-1] if ": " in cfg['label'] else cfg['label']
        lines.append((f"{section_no}. {short_label}:", True))
        for t in sorted(form_tickets, key=lambda x: x.get('id') or ''):
            lines.append((build_ticket_report_line(t, cfg), False))
    if not lines:
        lines = [("Không có công việc nào trong tuần này.", False)]
    return lines

def set_cell_lines(cell, lines):
    """Ghi nội dung mới vào 1 cell (list (text, bold)), mỗi tuple thành 1 paragraph riêng.
    Giữ định dạng (font) gốc của cell bằng cách tái sử dụng đoạn/run đầu tiên."""
    paragraphs = list(cell.paragraphs)
    for p in paragraphs[1:]:
        p._element.getparent().remove(p._element)
    base_p = cell.paragraphs[0]
    for r in list(base_p.runs):
        r._element.getparent().remove(r._element)
    if not lines:
        return
    first_text, first_bold = lines[0]
    run = base_p.add_run(first_text)
    run.bold = first_bold
    for text, bold in lines[1:]:
        p = cell.add_paragraph()
        r = p.add_run(text)
        r.bold = bold

def replace_paragraph_text(paragraph, new_text):
    """Thay toàn bộ nội dung 1 paragraph, giữ định dạng (font/bold/size) của run đầu tiên."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r._element.getparent().remove(r._element)

def update_title_paragraphs(doc, week, month, year, reporter):
    """Tìm và thay '(Tuần X tháng Y năm Z)' + 'Tên nhân viên báo cáo: ...' ở đầu file,
    bất kể nội dung cũ trong file mẫu là gì (tìm theo mẫu chữ, không theo vị trí cố định)."""
    week_pattern = re.compile(r"\(Tuần\s*\d+\s*tháng\s*\d+\s*năm\s*\d+\)")
    reporter_pattern = re.compile(r"^Tên nhân viên báo cáo\s*:")
    for p in doc.paragraphs:
        if week_pattern.search(p.text):
            replace_paragraph_text(p, f"(Tuần {week} tháng {month} năm {year})")
        elif reporter_pattern.match(p.text.strip()):
            replace_paragraph_text(p, f"Tên nhân viên báo cáo: {reporter}")

def find_main_content_table(doc):
    """Tìm bảng (kể cả lồng trong cell khác) có cột tiêu đề 'Các công việc thực hiện'."""
    def search(tables):
        for tbl in tables:
            if tbl.rows:
                header_texts = [c.text.strip() for c in tbl.rows[0].cells]
                if "Các công việc thực hiện" in header_texts:
                    return tbl
            for row in tbl.rows:
                for cell in row.cells:
                    found = search(cell.tables)
                    if found is not None:
                        return found
        return None
    return search(doc.tables)

def find_row_by_content(table, keyword, content_col_label="Nội dung"):
    """Tìm dòng trong bảng mà cột 'Nội dung' chứa keyword (vd: 'HỆ THỐNG BẢO HÀNH')."""
    header_cells = [c.text.strip() for c in table.rows[0].cells]
    col_idx = header_cells.index(content_col_label) if content_col_label in header_cells else 1
    for row in table.rows[1:]:
        if keyword in row.cells[col_idx].text:
            return row
    return None

def find_next_week_plan_table(doc):
    """Tìm bảng 'KẾ HOẠCH CÔNG VIỆC TUẦN SAU' — nhận diện qua cặp cột đặc trưng
    'Người thực hiện' + 'Ghi chú' (khác với bảng 'Tuần trước chưa giải quyết' dùng
    'Người giải quyết trực tiếp')."""
    def search(tables):
        for tbl in tables:
            if tbl.rows:
                header_texts = set()
                for r in tbl.rows[:2]:
                    for c in r.cells:
                        header_texts.add(c.text.strip())
                if "Người thực hiện" in header_texts and "Ghi chú" in header_texts:
                    return tbl
            for row in tbl.rows:
                for cell in row.cells:
                    found = search(cell.tables)
                    if found is not None:
                        return found
        return None
    return search(doc.tables)

def update_next_week_plan_notes(doc, note_text):
    """Điền khoảng ngày của TUẦN KẾ TIẾP vào cột 'Ghi chú' cho mọi dòng nội dung
    (1. ERP, 2. Hệ thống Bảo hành...) của bảng Kế hoạch công việc tuần sau."""
    table = find_next_week_plan_table(doc)
    if table is None:
        return False
    header_cells = [c.text.strip() for c in table.rows[0].cells]
    if "Ghi chú" not in header_cells:
        return False
    ghi_chu_idx = header_cells.index("Ghi chú")
    content_start = None
    for ri, row in enumerate(table.rows):
        if row.cells[0].text.strip().isdigit():
            content_start = ri
            break
    if content_start is None:
        return False
    for row in table.rows[content_start:]:
        set_cell_lines(row.cells[ghi_chu_idx], [(note_text, False)])
    return True

def generate_weekly_report_from_template(template_path, ctx):
    """Đổ dữ liệu vào ĐÚNG file mẫu công ty (giữ logo/layout/song ngữ/khung tự đánh giá...),
    chỉ thay: tiêu đề tuần, tên người báo cáo, nội dung + thời gian dòng HỆ THỐNG BẢO HÀNH,
    và cột Ghi chú (thời gian) của bảng Kế hoạch công việc tuần sau."""
    doc = Document(template_path)

    update_title_paragraphs(doc, ctx['week'], ctx['month'], ctx['year'], ctx['reporter'])

    table = find_main_content_table(doc)
    if table is None:
        raise ValueError("Không tìm thấy bảng có cột 'Các công việc thực hiện' trong file mẫu.")

    row = find_row_by_content(table, "HỆ THỐNG BẢO HÀNH")
    if row is None:
        raise ValueError("Không tìm thấy dòng 'HỆ THỐNG BẢO HÀNH' trong file mẫu.")

    header_cells = [c.text.strip() for c in table.rows[0].cells]
    content_idx = header_cells.index("Các công việc thực hiện")
    set_cell_lines(row.cells[content_idx], build_warranty_section_lines(ctx['warranty_tickets']))

    if "T/g giải quyết" in header_cells:
        time_idx = header_cells.index("T/g giải quyết")
        time_text = f"{ctx['start_date'].strftime('%d/%m')}-{ctx['end_date'].strftime('%d/%m')}/{ctx['end_date'].year}"
        set_cell_lines(row.cells[time_idx], [(time_text, False)])

    next_monday = ctx['start_date'] + datetime.timedelta(days=7)
    next_saturday = next_monday + datetime.timedelta(days=5)
    next_week_text = f"{next_monday.strftime('%d/%m')}-{next_saturday.strftime('%d/%m')}/{next_saturday.year}"
    update_next_week_plan_notes(doc, next_week_text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# -----------------
# GIAO DIỆN VIEW-ONLY TINH GỌN (GUEST VIEW)
# -----------------
if st.session_state.get("view_mode") == "only":
    # Ẩn Sidebar hoàn toàn bằng CSS
    st.markdown("""
    <style>
        [data-testid="stSidebar"] {
            display: none !important;
        }
        [data-testid="stSidebarCollapseButton"] {
            display: none !important;
        }
    </style>
    """, unsafe_allow_html=True)
    
    # Tìm kiếm Ticket trong db_tickets
    target_t = None
    for t in db_tickets:
        if t["id"] == url_ticket_id:
            target_t = t
            break
            
    if target_t:
        st.markdown(f'<div class="main-header">🔍 Tra cứu Ticket Chi tiết</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="sub-header">Giao diện xem nhanh thông tin yêu cầu bảo hành Gree IT</div>', unsafe_allow_html=True)
        
        # Thẻ thông tin tổng quan dạng card đẹp mắt
        with st.container(border=True):
            c_g1, c_g2, c_g3, c_g4 = st.columns([0.25, 0.25, 0.25, 0.25])
            with c_g1:
                st.markdown(f"**Mã Ticket:**\n`{target_t['id']}`")
            with c_g2:
                badge_html = status_badge(target_t['status'])
                st.markdown(f"**Trạng thái:**\n{badge_html}", unsafe_allow_html=True)
            with c_g3:
                st.markdown(f"**Người yêu cầu:**\n`{target_t['requester']}`")
            with c_g4:
                st.markdown(f"**Ngày khởi tạo:**\n`{format_ticket_date(target_t['created_at'])}`")
                
            st.markdown("---")
            st.markdown(f"**Tiêu đề yêu cầu:**\n### {target_t['subject']}")

        # Hàng chứa nút chuyển về hệ thống chính
        col_back_space, col_back_btn = st.columns([0.65, 0.35])
        with col_back_btn:
            if st.button("🏠 Mở trong hệ thống đầy đủ", type="primary", use_container_width=True):
                # Lưu trạng thái để chuyển về Full View
                st.session_state["view_mode"] = "full"
                st.session_state["active_ticket_id"] = target_t["id"]
                st.session_state["selected_page"] = "🛡️ Hệ thống Bảo hành"
                
                # Đồng bộ form_scope
                if target_t.get("form_type"):
                    disp_type = normalize_form_type(target_t.get("form_type"))
                    disp_config = FORM_TYPE_TO_CONFIG.get(disp_type)
                    if disp_config:
                        st.session_state["ticket_form_scope"] = disp_config["label"]
                        
                # Xóa tham số URL
                st.query_params.clear()
                st.rerun()

        # Dữ liệu đính kèm (nếu có)
        if target_t.get('form_data') and target_t.get('form_type'):
            display_form_type = normalize_form_type(target_t.get("form_type"))
            display_form_name = FORM_TYPE_TO_CONFIG.get(display_form_type, {}).get("label", display_form_type)
            st.markdown("---")
            st.markdown(f"#### 📋 Dữ liệu đính kèm: **{display_form_name}**")
            
            # Hiển thị dạng bảng (nếu thuộc cấu hình Form) hoặc JSON
            display_config = FORM_TYPE_TO_CONFIG.get(display_form_type)
            if display_config:
                # Dựng 1 DataFrame gồm 1 hàng duy nhất cho ticket này
                rows_data = build_form_rows([target_t], display_config)
                if rows_data:
                    st.dataframe(pd.DataFrame(rows_data), use_container_width=True)
            else:
                try:
                    f_data = parse_form_data(target_t)
                    st.json(f_data)
                except Exception as e:
                    st.error(f"Lỗi hiển thị dữ liệu Form: {e}")

        # Lịch sử hội thoại (Chat-log)
        st.markdown("---")
        st.markdown("#### 💬 Lịch sử trao đổi & Xử lý (Chat-log)")
        
        # Vẽ các tin nhắn
        for m in target_t['msgs']:
            is_admin = m['user'].startswith("Admin")
            # Trong chế độ guest, ẩn log nội bộ để bảo mật!
            if m['type'] == "internal":
                continue
            with st.chat_message("assistant" if is_admin else "user"):
                label = "🔒 Nội bộ" if m['type'] == "internal" else "🌐 Công khai"
                st.write(f"**{m['user']}** ({label})")
                st.write(m['msg'])
                # Hiển thị ngày và giờ cùng nhau
                st.caption(f"{m.get('date', '')} • {m.get('time', '')}")

        # Gửi log mới (chỉ cho phép Công khai trong chế độ Guest)
        st.divider()
        st.markdown("##### 📩 Gửi phản hồi mới")
        guest_name = st.text_input("Tên của bạn", value=target_t['requester'], key="guest_msg_name")
        guest_msg = st.text_area("Nội dung phản hồi...", height=80, key="guest_msg_val")
        
        if st.button("📩 Gửi phản hồi", use_container_width=True, key="btn_guest_send"):
            if guest_msg and guest_name:
                db.add_ticket_message(target_t['id'], guest_name, guest_msg, "public")
                st.success("Đã gửi phản hồi thành công!")
                st.rerun()
            else:
                st.warning("Vui lòng nhập đầy đủ tên và nội dung phản hồi.")
    else:
        st.error(f"❌ Không tìm thấy Ticket với mã yêu cầu `{url_ticket_id}` hoặc yêu cầu đã bị xóa vĩnh viễn khỏi hệ thống.")
        if st.button("🏠 Quay lại trang chủ hệ thống", type="primary"):
            st.session_state["view_mode"] = "full"
            st.query_params.clear()
            st.rerun()
            
    # Dừng chạy code phía dưới để giữ giao diện tinh gọn
    st.stop()

# -----------------
# THANH ĐIỀU HƯỚNG (SIDEBAR)
# -----------------
# Khởi tạo trạng thái sitemap page mặc định
if "selected_page" not in st.session_state:
    st.session_state["selected_page"] = "🏠 Trang chủ"

with st.sidebar:
    st.title(Config.APP_NAME)
    st.markdown(f"**Owner:** {Config.OWNER}")
    st.markdown(f"**Domain:** [{Config.BASE_DOMAIN}{Config.SUB_PATH}](https://{Config.BASE_DOMAIN}{Config.SUB_PATH})")
    st.markdown("---")
    
    menu = [
        "🏠 Trang chủ", 
        "🛡️ Hệ thống Bảo hành", 
        "🏭 Quản trị ERP", 
        "📱 Gree App Support", 
        #"🔍 Tra cứu An Gia", 
        #"📦 Quản lý XNK", 
        #"🔐 Security Audit", 
        #"🌐 Website", 
        "📈 Báo cáo tuần của Nam", 
        "⚙️ Cài đặt"
    ]
    
    # Tìm chỉ số trang hiện tại trong menu để làm default index
    default_page_idx = 0
    if st.session_state["selected_page"] in menu:
        default_page_idx = menu.index(st.session_state["selected_page"])
        
    page = st.radio("SITEMAP HỆ THỐNG", menu, index=default_page_idx)
    st.session_state["selected_page"] = page
    
    st.markdown("---")
    st.caption(f"Admin: {Config.OWNER} | 📅 {datetime.date.today()}")

# -----------------
# GIAO DIỆN CHÍNH
# -----------------

if page == "🏠 Trang chủ":
    st.markdown(f'<div class="main-header">Welcome to {Config.APP_NAME}</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">Dashboard Tổng quát</div>', unsafe_allow_html=True)
    
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Tổng số Ticket", str(len(db_tickets)), "")
    c2.metric("Đang xử lý", len([t for t in db_tickets if t['status'] not in ['Hoàn thành', 'Từ chối']]))
    c3.metric("Hoàn thành (Tasks)", len(db_tasks))
    c4.metric("NCC Issue", "1", "Gấp")

    st.divider()
    cl, cr = st.columns(2)
    with cl:
        st.subheader("📅 Lịch trình & Ghi chú")
        st.write("- [ ] Họp ERP chiều thứ 2 (05/05)")
        st.text_area("Ghi chú nhanh", placeholder="Nhập mã lỗi hoặc ID cần lưu ý...")
    with cr:
        st.subheader("🔗 Truy cập nhanh")
        st.button("Hệ thống ERP", use_container_width=True)
        st.button("Gree App Admin", use_container_width=True)

elif page == "🛡️ Hệ thống Bảo hành":
    st.markdown('<div class="main-header">🛡️ Hệ thống Warranty - Bảo Hành</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">Quản lý và hỗ trợ xử lý lỗi hệ thống</div>', unsafe_allow_html=True)

    form_scope_options = ["Tất cả Ticket"] + [f["label"] for f in FORM_CONFIGS]
    selected_scope = st.selectbox(
        "Form đang xem",
        form_scope_options,
        key="ticket_form_scope",
    )
    selected_form_type = FORM_LABEL_TO_TYPE.get(selected_scope)
    selected_config = FORM_TYPE_TO_CONFIG.get(selected_form_type)
    scoped_tickets = tickets_for_form(db_tickets, selected_form_type) if selected_form_type else db_tickets
    
    btn_col, hint_col = st.columns([0.28, 0.72])
    with btn_col:
        if st.button("➕ Tạo Ticket mới", use_container_width=True, disabled=selected_config is None):
            st.session_state["show_create_ticket"] = not st.session_state.get("show_create_ticket", False)
    with hint_col:
        if selected_config:
            st.caption(f"Ticket mới sẽ được tạo cho: {selected_scope}")
        else:
            st.caption("Chọn một Form cụ thể để tạo ticket mới.")

    if selected_config and st.session_state.get("show_create_ticket", False):
        form_choice = selected_scope
        form_choice_type = selected_form_type
        with st.container(border=True):
            st.subheader(f"➕ Tạo Ticket mới - {selected_scope}")
        
            t_user = st.text_input("Người yêu cầu", placeholder="Ví dụ: Phượng BH RAC")
        
            form_data = {}
            if form_choice_type == "Khai_Bao_Model_Bao_Hanh":
                left_col, right_col = st.columns(2)
                with left_col:
                    form_data['model_name'] = st.text_input("Tên Model", placeholder="Ví dụ: GWC12PB-K3D0P4")
                    form_data['cost_type'] = st.selectbox("Loại chi phí", ["CAC", "GD - RAC", "ĐH - RAC", "RAC - CT"])
                    form_data['product_type'] = st.selectbox("Loại sản phẩm", [
                        "BĐT - Bếp điện từ", 
                        "Điều hòa dân dụng, treo tường tú đứng dưới 10Hp (RAC)", 
                        "MHA - Máy hút ẩm", 
                        "MLKK - Máy lọc không khí", 
                        "MLM - Máy làm mát bằng hơi nước", 
                        "Multi, máy âm trần, âm trần nối ống gió, áp trần (CAC)", 
                        "NAS - Nồi áp suất", 
                        "NCĐ - nồi cơm điện", 
                        "Ống gió lớn, Tú đứng lớn, máy lạnh chính xác", 
                        "QĐ - Quạt điện", 
                        "VRV/VRF(CAC)", 
                        "Water cooled packge Chiller"
                    ])
                with right_col:
                    form_data['capacity_range'] = st.selectbox("Công suất", [
                        "9.000 - 18.000 Btu", "9.000 - 36.000 Btu", "24.000 - 36.000 Btu", 
                        "36.000 - 42.000 Btu", "42.000 - 60.000 Btu", "100.000 - 200.000 Btu", 
                        "10Hp - 24Hp", "10Hp - 60Hp"
                    ])
                    form_data['note'] = st.text_area("Ghi chú", height=96, placeholder="Thông tin bổ sung nếu có")
            elif form_choice_type == "Khai_Bao_Ma_Linh_Kien":
                form_data['part_code'] = st.text_input("Mã linh kiện")
                form_data['part_name_en'] = st.text_input("Tên linh kiện (EN)")
                form_data['part_name_vi'] = st.text_input("Tên linh kiện (VI)")
                form_data['description'] = st.text_area("Mô tả")
            elif form_choice_type == "Yeu_Cau_Dieu_Chinh_Ton_Kho":
                left_col, right_col = st.columns(2)
                with left_col:
                    form_data['warehouse'] = st.text_input("Kho/Trạm", placeholder="Ví dụ: Hưng Yên - RAC")
                    form_data['part_code'] = st.text_input("Mã linh kiện", placeholder="Ví dụ: GMC42S6I1")
                    form_data['part_nature'] = st.selectbox("Tính chất linh kiện", ["", "Linh kiện mượn", "Linh kiện mua"])
                    form_data['evidence_image_url'] = st.text_input("Link ảnh bằng chứng", placeholder="https://prnt.sc/...")
                with right_col:
                    form_data['export_voucher'] = st.text_input("Phiếu xuất", placeholder="Để trống nếu chưa có")
                    form_data['adjusted_quantity'] = st.number_input("Số lượng điều chỉnh", min_value=0, value=1)
                    form_data['note'] = st.text_area(
                        "Ghi chú xử lý",
                        height=140,
                        placeholder="Mô tả lệch tồn, hướng xử lý, lưu ý cho kho/trạm",
                    )
            elif form_choice_type == "Dang_Ky_Tram_Bao_Hanh_Moi":
                left_col, right_col = st.columns(2)
                with left_col:
                    st.markdown("**Thông tin công ty**")
                    company_name = st.text_input("Tên công ty / Trạm", placeholder="Ví dụ: CÔNG TY TNHH TM DV A.T.P")
                    tax_code = st.text_input("Mã số thuế", placeholder="Ví dụ: 0301841221")
                    tax_address = st.text_area("Địa chỉ thuế", height=90)
                    postal_address = st.text_area("Địa chỉ nhận thư", height=90)
                    email = st.text_input("Email", placeholder="email@domain.com")
                    phone = st.text_input("Điện thoại", placeholder="Ví dụ: 0916318948")
                    system_username = st.text_input("User hệ thống", placeholder="Ví dụ: a.t.p-hochiminh")
                with right_col:
                    st.markdown("**Tài khoản & cấu trúc liên quan**")
                    bank_account = st.text_input("Số tài khoản")
                    bank_account_name = st.text_input("Tên tài khoản ngân hàng")
                    bank_name = st.text_input("Ngân hàng")
                    technicians_raw = st.text_area(
                        "Danh sách KTV",
                        height=120,
                        placeholder="Mỗi dòng: Tên KTV | SĐT | Username",
                    )
                    warehouses_raw = st.text_area(
                        "Danh sách kho liên kết",
                        height=120,
                        placeholder="Mỗi dòng: Loại kho | Tên kho | Email",
                    )
                form_data = {
                    "company_info": {
                        "name": company_name,
                        "tax_code": tax_code,
                        "tax_address": tax_address,
                        "postal_address": postal_address,
                        "email": email,
                        "phone": phone,
                        "bank_account": bank_account,
                        "bank_account_name": bank_account_name,
                        "bank_name": bank_name,
                        "system_username": system_username,
                    },
                    "technicians": parse_pipe_rows(technicians_raw, ["name", "phone", "username"]),
                    "associated_warehouses": parse_pipe_rows(warehouses_raw, ["type", "name", "email"]),
                }
            elif form_choice_type == "Dang_Ky_Tai_Khoan_User_Noi_Bo":
                left_col, right_col = st.columns(2)
                with left_col:
                    form_data['full_name'] = st.text_input("Họ tên", placeholder="Ví dụ: Võ Thanh Tùng")
                    form_data['phone'] = st.text_input("Điện thoại", placeholder="Ví dụ: 0797 704 205")
                    form_data['company_email'] = st.text_input("Email Gree", placeholder="name@gree.com.vn")
                    form_data['user_group'] = st.selectbox(
                        "Nhóm user",
                        [
                            "Giám đốc BH",
                            "Trưởng phòng BH",
                            "KTV Gree",
                            "Trưởng phòng XNK",
                            "Chuyên viên XNK",
                            "Phòng Thương Mại CAC",
                            "Phòng Quản Trị",
                            "Chuyên viên Kế Hoạch",
                            "Trưởng phòng Kế hoạch",
                            "Trưởng Phòng Call Center",
                            "Chuyên Viên Call Center",
                        ],
                    )
                with right_col:
                    form_data['call_center_line'] = st.text_input("Line Call Center", placeholder="Để trống nếu không áp dụng")
                    form_data['main_link'] = st.text_input("Link chính", value="https://warranty.gree.com.vn/")
                    test_username = st.text_input("Username test", value="test")
                    test_password = st.text_input("Password test", value="123456")
                    test_link = st.text_input("Link test", value="http://gree.baohanhso.net/")
                    form_data['test_account'] = {
                        "username": test_username,
                        "password": test_password,
                        "test_link": test_link,
                    }
            elif form_choice_type == "Khai_Bao_Model_Ho_So_May":
                form_data['model_name'] = st.text_input("Tên Model")
                form_data['machine_type'] = st.selectbox("Loại máy", ["Dàn nóng", "Dàn lạnh"])
                form_data['warranty_months_machine'] = st.number_input("T/g BH Máy (Tháng)", min_value=0, value=24)
                form_data['warranty_months_compressor'] = st.number_input("T/g BH Block (Tháng)", min_value=0, value=36)
            elif form_choice_type == "Admin_Web_Noted_Log_Ho_Tro":
                left_col, right_col = st.columns(2)
                with left_col:
                    form_data['case_code'] = st.text_input("Mã ca / Mã chứng từ", placeholder="Ví dụ: MBTH2025080234")
                    form_data['it_assignee'] = st.text_input("Người xử lý IT", value="IT - Nam Lê")
                    form_data['completion_date'] = st.date_input("Ngày hoàn thành", value=datetime.date.today()).isoformat()
                    form_data['evidence_link'] = st.text_input("Link bằng chứng", placeholder="https://prnt.sc/...")
                with right_col:
                    form_data['request_detail'] = st.text_area(
                        "Nội dung yêu cầu",
                        height=120,
                        placeholder="Nhập nội dung cần hỗ trợ/can thiệp dữ liệu",
                    )
                    form_data['internal_action_log'] = st.text_area(
                        "Log xử lý nội bộ",
                        height=120,
                        placeholder="Nhập thao tác đã xử lý, kết quả, lưu ý nội bộ",
                    )
            elif form_choice_type == "Import_Bang_Gia_Linh_Kien":
                form_data['part_code'] = st.text_input("Mã linh kiện")
                form_data['part_name_vi'] = st.text_input("Tên linh kiện (VI)")
                form_data['part_name_en'] = st.text_input("Tên linh kiện (EN)")
                form_data['model_group'] = st.text_input("Nhóm Model")
                form_data['product_type'] = st.text_input("Loại sản phẩm")
                form_data['unit_type'] = st.text_input("Loại Unit (Indoor/Outdoor)")
                form_data['price_vat'] = st.number_input("Giá có VAT", min_value=0, value=0)
                form_data['price_no_vat'] = st.number_input("Giá chưa VAT", min_value=0, value=0)
                form_data['classification'] = st.text_input("Phân loại (Vd: CAC)")
                form_data['discount_rate'] = st.text_input("Tỷ lệ chiết khấu (%)")
                form_data['sla_bonus_rate'] = st.text_input("Thưởng SLA (%)")
            elif form_choice_type == "Khai_Bao_Danh_Muc_Chi_Phi":
                form_data['cost_code'] = st.text_input("Mã chi phí")
                form_data['content'] = st.text_input("Nội dung")
                form_data['unit_price'] = st.number_input("Đơn giá", min_value=0, value=0)
                form_data['capacity'] = st.text_input("Công suất (Tùy chọn)")
                form_data['cost_type'] = st.text_input("Loại chi phí")
                form_data['product_group'] = st.text_input("Nhóm sản phẩm")
            elif form_choice_type == "Yeu_Cau_Tong_Hop_Khac":
                form_data['request_title'] = st.text_input(
                    "Tiêu đề yêu cầu",
                    placeholder="Ví dụ: Hỗ trợ kiểm tra số liệu tồn kho khu vực Hà Nội",
                )
                related_form_options = [
                    f["label"] for f in FORM_CONFIGS if f["type"] != "Yeu_Cau_Tong_Hop_Khac"
                ] + ["Khác (Không thuộc 9 Form trên)"]
                form_data['related_form'] = st.selectbox(
                    "Yêu cầu này liên quan đến Form nào?",
                    related_form_options,
                )
                form_data['request_detail'] = st.text_area(
                    "Nội dung yêu cầu",
                    height=160,
                    placeholder="Mô tả chi tiết yêu cầu cần hỗ trợ/xử lý...",
                )

            create_col, close_col = st.columns(2)
            if create_col.button("Khởi tạo mã Ticket", use_container_width=True):
                form5_missing = missing_form5_required_fields(form_data) if form_choice_type == "Dang_Ky_Tai_Khoan_User_Noi_Bo" else []
                form10_missing = []
                if form_choice_type == "Yeu_Cau_Tong_Hop_Khac" and not str(form_data.get('request_detail') or "").strip():
                    form10_missing.append("Nội dung yêu cầu")

                if form5_missing:
                    st.warning("Vui lòng nhập đủ thông tin Form 5: " + ", ".join(form5_missing))
                elif form10_missing:
                    st.warning("Vui lòng nhập đủ thông tin: " + ", ".join(form10_missing))
                elif t_user:
                    new_id = gen_ticket_id()
                    form_type = form_choice_type
                    status = "Hoàn thành" if form_choice_type == "Admin_Web_Noted_Log_Ho_Tro" else "Mới tạo"
                    subject = form_choice
                    if form_choice_type == "Yeu_Cau_Tong_Hop_Khac" and str(form_data.get('request_title') or "").strip():
                        subject = form_data['request_title']
                    
                    if db.create_ticket(new_id, subject, t_user, form_type, form_data, status):
                        st.success(f"Đã tạo Ticket: {new_id}")
                        st.session_state["show_create_ticket"] = False
                        st.rerun()
                    else:
                        st.error("Lỗi khi kết nối CSDL")
                else:
                    st.warning("Vui lòng nhập Người yêu cầu")
            if close_col.button("Đóng form", use_container_width=True):
                st.session_state["show_create_ticket"] = False
                st.rerun()

    st.markdown("---")

    st.subheader("📊 Dữ liệu danh mục đã cập nhật hệ thống")
    
    # Xác định loại Form cần hiển thị dựa trên selected_scope ở đầu trang
    master_form_options = [f["label"] for f in FORM_CONFIGS]
    
    if selected_scope in master_form_options:
        # Nếu đang chọn một Form cụ thể trên đầu trang
        target_form_label = selected_scope
        is_default_view = False
    else:
        # Nếu đang chọn "Tất cả Ticket", mặc định hiển thị Form 1
        target_form_label = master_form_options[0]  # Form 1
        is_default_view = True
        st.caption("ℹ️ *Đang hiển thị mặc định danh sách Form 1 (Model phân loại chi phí). Chọn một Form cụ thể ở ô 'Form đang xem' trên đầu trang để đồng bộ bảng dữ liệu này.*")
    
    master_form_type = FORM_LABEL_TO_TYPE.get(target_form_label)
    master_config = FORM_TYPE_TO_CONFIG.get(master_form_type)
    
    if master_config:
        # Lấy tickets của form đó
        form_tickets = tickets_for_form(db_tickets, master_form_type)
        master_status_filter = st.selectbox(
            "🔍 Lọc trạng thái bản ghi",
            ["Tất cả"] + STATUS_LIST,
            key=f"master_filter_{master_form_type}",
        )
        master_tickets = form_tickets if master_status_filter == "Tất cả" else [
            t for t in form_tickets if t["status"] == master_status_filter
        ]
        master_rows = build_form_rows(master_tickets, master_config)
        if master_rows:
            st.dataframe(pd.DataFrame(master_rows), use_container_width=True)
        else:
            st.info(f"Chưa có dữ liệu hoặc không có bản ghi nào ở trạng thái '{master_status_filter}' cho form này.")
    else:
        st.info("Không tìm thấy cấu hình hiển thị dữ liệu cho Form này.")

    st.divider()
    
    col_list, col_detail = st.columns([0.4, 0.6])
    
    with col_list:
        st.subheader("📋 Danh sách Ticket")
        ticket_filter = st.selectbox("🔍 Lọc trạng thái", ["Tất cả"] + STATUS_LIST, key="ticket_filter")
        filtered_tickets = scoped_tickets if ticket_filter == "Tất cả" else [
            t for t in scoped_tickets if t["status"] == ticket_filter
        ]
        
        if not filtered_tickets:
            st.info("Không có ticket nào trong phạm vi đang chọn.")

        for t in filtered_tickets:
            with st.container(border=True):
                badge_html = status_badge(t['status'])
                st.markdown(f"**{t['id']}** &nbsp; {badge_html}", unsafe_allow_html=True)
                st.caption(f"Từ: {t['requester']} - {t['subject']}")
                if st.button("Xem chi tiết", key=f"btn_{t['id']}"):
                    st.session_state['active_ticket_id'] = t["id"]

    with col_detail:
        st.subheader("💬 Luồng xử lý (Chat-log)")
        active_ticket_id = st.session_state.get("active_ticket_id")
        active_candidates = [t for t in scoped_tickets if t["id"] == active_ticket_id]
        if active_ticket_id and not active_candidates:
            st.info("Ticket đang chọn không thuộc Form/phạm vi hiện tại. Hãy chọn ticket bên trái.")
        elif active_candidates:
            current_t = active_candidates[0]
            
            # Đường dẫn chia sẻ ticket nhanh dành cho BA/Admin
            st.markdown("🔗 **Liên kết chia sẻ ticket nhanh (BA-Share):**")
            
            ticket_json = json.dumps(current_t['id'])
            
            share_html = (
                "<div style='font-family: Inter, sans-serif; font-size: 0.95rem; color: #111;'>"
                "  <div style='margin-bottom: 0.5rem;'>"
                "    <span style='font-weight: 600;'>Link hiện tại:</span>"
                "  </div>"
                "  <div style='display: flex; align-items: center; gap: 0.75rem;'>"
                "    <code id='share_url_text' style='display: inline-block; padding: 0.6rem 0.8rem; background: #f7f7f8; border-radius: 0.75rem; color: #0f172a; overflow-x: auto; white-space: nowrap; max-width: 100%;'>Đang tạo...</code>"
                "    <button id='copy_button' style='border: none; padding: 0.6rem 1rem; border-radius: 0.75rem; background: #2563eb; color: white; cursor: pointer;'>Copy</button>"
                "  </div>"
                "  <div style='margin-top: 0.5rem;'>"
                "    <a id='share_url_link' href='#' target='_blank' rel='noreferrer' style='color: #2563eb; text-decoration: none; font-weight: 600;'></a>"
                "  </div>"
                "</div>"
                "<script>"
                "const ticketId = " + ticket_json + ";"
                "try {"
                "  const origin = window.parent.location.origin;"
                "  const shareUrl = origin + '/?ticket=' + ticketId;"
                "  const textEl = document.getElementById('share_url_text');"
                "  const linkEl = document.getElementById('share_url_link');"
                "  const copyButton = document.getElementById('copy_button');"
                "  if (textEl) { textEl.textContent = shareUrl; }"
                "  if (linkEl) { linkEl.href = shareUrl; linkEl.textContent = 'Mở liên kết'; }"
                "  if (copyButton) { copyButton.addEventListener('click', () => { navigator.clipboard.writeText(shareUrl).then(() => { copyButton.textContent = 'Copied'; setTimeout(() => { copyButton.textContent = 'Copy'; }, 1500); }); }); }"
                "} catch(e) {"
                "  console.error('Share URL error:', e);"
                "}"
                "</script>"
            )
            st.components.v1.html(share_html, height=140)
            
            # Nếu có dữ liệu form thì in ra JSON format đẹp mắt
            if current_t.get('form_data') and current_t.get('form_type'):
                display_form_type = normalize_form_type(current_t.get("form_type"))
                display_form_name = FORM_TYPE_TO_CONFIG.get(display_form_type, {}).get("label", display_form_type)
                st.info(f"📋 Dữ liệu đính kèm: **{display_form_name}**")
                try:
                    f_data = parse_form_data(current_t)
                    st.json(f_data)
                except Exception as e:
                    st.error(f"Lỗi hiển thị dữ liệu Form: {e}")
                    
            # Hiển thị hội thoại
            for m in current_t['msgs']:
                is_admin = m['user'].startswith("Admin")
                with st.chat_message("assistant" if is_admin else "user"):
                    label = "🔒 Nội bộ" if m['type'] == "internal" else "🌐 Công khai"
                    st.write(f"**{m['user']}** ({label})")
                    st.write(m['msg'])
                    # Hiển thị ngày và giờ cùng nhau
                    st.caption(f"{m.get('date', '')} • {m.get('time', '')}")

            # Nhập Log mới
            st.divider()
            log_msg = st.text_area("Nhập nội dung xử lý (Log)...", height=120, placeholder="Bạn có thể nhập nhiều dòng nội dung log xử lý tại đây...")
            log_type = st.radio("Loại log:", ["Công khai", "Nội bộ (Chỉ Admin)"], horizontal=True)
            
            c_btn1, c_btn2 = st.columns(2)
            if c_btn1.button("📩 Gửi Log", use_container_width=True):
                if log_msg:
                    mtype = "internal" if "Nội bộ" in log_type else "public"
                    db.add_ticket_message(current_t['id'], "Admin Nam", log_msg, mtype)
                    st.rerun()
                else:
                    st.warning("Vui lòng nhập nội dung")
                
            # Dropdown chuyển trạng thái
            st.divider()
            st.markdown("**🔄 Chuyển trạng thái Ticket:**")
            current_status = current_t['status']
            new_status = st.selectbox(
                "Chọn trạng thái mới",
                STATUS_LIST,
                index=STATUS_LIST.index(current_status) if current_status in STATUS_LIST else 0,
                key="status_change"
            )
            col_s1, col_s2 = st.columns(2)
            if col_s1.button("🔄 Cập nhật trạng thái", use_container_width=True):
                if new_status != current_status:
                    if new_status == "Hoàn thành":
                        action_log = log_msg if log_msg else "Admin đã đánh dấu hoàn thành"
                        db.complete_ticket(current_t['id'], current_t['subject'], action_log)
                        st.balloons()
                    else:
                        db.update_ticket_status(current_t['id'], new_status)
                    st.rerun()
                else:
                    st.info("Trạng thái không thay đổi.")
            if col_s2.button("❌ Từ chối Ticket", use_container_width=True):
                if current_status != "Từ chối":
                    db.update_ticket_status(current_t['id'], "Từ chối")
                    st.rerun()
                else:
                    st.info("Ticket này đã bị từ chối rồi.")

            # Khu vực Quản trị Admin - Sửa và Xóa Ticket
            st.divider()
            with st.expander("🛠️ Quản trị Admin - Chỉnh sửa / Xóa Ticket"):
                st.markdown("##### ✏️ Chỉnh sửa thông tin Ticket")
                edit_subject = st.text_input("Tiêu đề Ticket", value=current_t['subject'], key=f"edit_subj_{current_t['id']}")
                edit_requester = st.text_input("Người yêu cầu", value=current_t['requester'], key=f"edit_req_{current_t['id']}")
                
                # Chỉnh sửa form_data JSON
                edit_form_json = ""
                has_form = False
                if current_t.get('form_data') and current_t.get('form_type'):
                    has_form = True
                    current_form_data = parse_form_data(current_t)
                    edit_form_json = st.text_area(
                        "Dữ liệu đính kèm (JSON)", 
                        value=json.dumps(current_form_data, indent=4, ensure_ascii=False),
                        height=200,
                        key=f"edit_json_{current_t['id']}"
                    )
                
                c_edit1, c_edit2 = st.columns(2)
                if c_edit1.button("💾 Lưu Thay Đổi", type="primary", use_container_width=True, key=f"btn_save_{current_t['id']}"):
                    parsed_form_data = None
                    json_ok = True
                    if has_form:
                        try:
                            parsed_form_data = json.loads(edit_form_json)
                        except Exception as e:
                            st.error(f"Dữ liệu JSON không hợp lệ: {e}")
                            json_ok = False
                    
                    if json_ok:
                        if db.update_ticket_data(current_t['id'], edit_subject, edit_requester, parsed_form_data):
                            st.success("Đã cập nhật thông tin ticket thành công!")
                            st.rerun()
                        else:
                            st.error("Lỗi khi cập nhật CSDL")
                            
                st.markdown("---")
                st.markdown("##### 🗑️ Xóa Ticket vĩnh viễn")
                confirm_delete = st.checkbox("Tôi xác nhận muốn xóa vĩnh viễn ticket này khỏi hệ thống.", key=f"conf_del_{current_t['id']}")
                if st.button("🗑️ Thực hiện Xóa Ticket", type="primary", disabled=not confirm_delete, use_container_width=True, key=f"btn_del_{current_t['id']}"):
                    if db.delete_ticket(current_t['id']):
                        st.success("Đã xóa ticket thành công!")
                        if 'active_ticket_id' in st.session_state:
                            del st.session_state['active_ticket_id']
                        st.rerun()
                    else:
                        st.error("Lỗi khi xóa ticket khỏi CSDL")
        else:
            st.info("Chọn một ticket bên trái để xem chi tiết.")

elif page == "📈 Báo cáo tuần của Nam":
    st.markdown('<div class="main-header">📊 Báo cáo công việc theo tuần</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">Tuần tính theo tháng (Tuần 1 → Tuần 5 mỗi tháng, reset khi sang tháng mới)</div>', unsafe_allow_html=True)

    today = datetime.date.today()
    wc1, wc2, wc3, wc4 = st.columns([1, 1, 1, 2])
    with wc1:
        report_year = st.selectbox("Năm", list(range(today.year - 1, today.year + 2)), index=1)
    with wc2:
        report_month = st.selectbox("Tháng", list(range(1, 13)), index=today.month - 1, format_func=lambda m: f"Tháng {m}")
    with wc3:
        max_w = max_weeks_in_month(report_year, report_month)
        default_week = week_of_month(today) if (report_year == today.year and report_month == today.month) else 1
        default_week = min(default_week, max_w)
        report_week = st.selectbox("Tuần", list(range(1, max_w + 1)), index=default_week - 1, format_func=lambda w: f"Tuần {w}")
    with wc4:
        reporter_name = st.text_input("Người báo cáo", value=Config.OWNER)

    start_date, end_date = week_date_range(report_year, report_month, report_week)
    st.caption(
        f"📅 Khoảng thời gian: **{start_date.strftime('%d/%m/%Y')} → {end_date.strftime('%d/%m/%Y')}** "
        f"(Tuần {report_week} tháng {report_month}/{report_year})"
    )

    # Ticket được TẠO trong tuần đang chọn
    week_new_tickets = [
        t for t in db_tickets
        if to_date(t.get('created_at')) and start_date <= to_date(t.get('created_at')) <= end_date
    ]
    # Công việc HOÀN THÀNH trong tuần đang chọn — dùng tasks.created_at (ghi nhận đúng thời điểm
    # đánh dấu Hoàn thành, không phải ngày tạo ticket), join lại với tickets qua ticket_id để lấy đủ thông tin.
    tickets_by_id = {t['id']: t for t in db_tickets}
    week_completed_tasks = [
        task for task in db_tasks
        if to_date(task.get('created_at')) and start_date <= to_date(task.get('created_at')) <= end_date
    ]

    st.divider()
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Ticket mới trong tuần", len(week_new_tickets))
    m2.metric("Hoàn thành trong tuần", len(week_completed_tasks))
    m3.metric("Đang xử lý/Chờ xử lý", len([t for t in week_new_tickets if t['status'] in ['Đang xử lý', 'Chờ xử lý']]))
    m4.metric("Từ chối", len([t for t in week_new_tickets if t['status'] == 'Từ chối']))

    summary_rows = []
    for cfg in FORM_CONFIGS:
        cnt = len([t for t in week_new_tickets if normalize_form_type(t.get('form_type')) == cfg['type']])
        if cnt:
            summary_rows.append({"Form": cfg['label'], "Số lượng": cnt})

    if summary_rows:
        st.subheader("📊 Ticket mới theo Form")
        st.dataframe(pd.DataFrame(summary_rows), use_container_width=True, hide_index=True)

    st.subheader("✅ Danh sách công việc đã hoàn thành trong tuần")
    completed_rows = []
    for task in week_completed_tasks:
        ref_ticket = tickets_by_id.get(task.get('ticket_id'))
        if ref_ticket:
            f_type = normalize_form_type(ref_ticket.get('form_type'))
            f_label = FORM_TYPE_TO_CONFIG.get(f_type, {}).get('label', f_type or '—')
            completed_rows.append({
                "id": ref_ticket.get("id", ""),
                "form_label": f_label,
                "requester": ref_ticket.get("requester", ""),
                "subject": ref_ticket.get("subject", task.get("content", "")),
            })
        else:
            # Log cũ chưa có ticket_id (ghi trước khi nâng cấp), hoặc ticket gốc đã bị xóa
            completed_rows.append({
                "id": "—",
                "form_label": "—",
                "requester": "—",
                "subject": task.get("content", task.get("action", "")),
            })
    if completed_rows:
        st.dataframe(
            pd.DataFrame(completed_rows).rename(columns={
                "id": "Mã Ticket", "form_label": "Loại Form", "requester": "Người y/c", "subject": "Nội dung"
            }),
            use_container_width=True, hide_index=True,
        )
    else:
        st.info("Chưa có công việc nào hoàn thành trong tuần này.")

    st.caption(
        "ℹ️ *Số liệu 'Hoàn thành trong tuần' lấy theo thời điểm thực tế đánh dấu Hoàn thành (bảng tasks), không phải "
        "ngày tạo ticket. Các dòng hiển thị 'Mã Ticket: —' là log được ghi trước khi cập nhật cột `ticket_id` (17/6/2026), "
        "nên chưa liên kết lại được với ticket gốc.*"
    )

    st.divider()
    st.subheader("📝 Bổ sung nội dung báo cáo (nhập tay)")
    manual_tasks = st.text_area("Công việc khác ngoài hệ thống ticket", placeholder="Những việc đã làm nhưng chưa tạo ticket...")
    issues = st.text_area("Vấn đề gặp phải / Rủi ro cần lưu ý", placeholder="Nhập các vấn đề nếu có...")

    # Tập hợp ticket "liên quan tuần này" để đổ vào mục HỆ THỐNG BẢO HÀNH của mẫu công ty:
    # = ticket mới tạo trong tuần ∪ ticket được đánh dấu hoàn thành trong tuần (dù tạo trước đó)
    warranty_ticket_ids = set(t['id'] for t in week_new_tickets)
    for task in week_completed_tasks:
        if task.get('ticket_id'):
            warranty_ticket_ids.add(task['ticket_id'])
    warranty_tickets = [t for t in db_tickets if t['id'] in warranty_ticket_ids]

    st.divider()
    if not DOCX_AVAILABLE:
        st.error("Chưa cài thư viện `python-docx`. Vui lòng chạy `pip install python-docx` rồi khởi động lại app để dùng tính năng xuất Word.")
    else:
        report_ctx = {
            "year": report_year, "month": report_month, "week": report_week,
            "reporter": reporter_name, "start_date": start_date, "end_date": end_date,
            "metrics": [
                ("Ticket mới trong tuần", len(week_new_tickets)),
                ("Hoàn thành trong tuần", len(week_completed_tasks)),
                ("Đang xử lý/Chờ xử lý", len([t for t in week_new_tickets if t['status'] in ['Đang xử lý', 'Chờ xử lý']])),
                ("Từ chối", len([t for t in week_new_tickets if t['status'] == 'Từ chối'])),
            ],
            "summary_rows": summary_rows,
            "completed_rows": completed_rows,
            "manual_tasks": manual_tasks,
            "issues": issues,
            "warranty_tickets": warranty_tickets,
        }

        exp_col1, exp_col2 = st.columns(2)
        with exp_col1:
            st.caption("Dùng đúng file mẫu công ty (logo, song ngữ, khung tự đánh giá...) — chỉ đổ dữ liệu vào mục Hệ thống Bảo hành.")
            if st.button("📄 Xuất theo Mẫu Công ty (.docx)", type="primary", use_container_width=True):
                if not os.path.exists(TEMPLATE_PATH):
                    st.error(
                        f"Không tìm thấy file mẫu tại `{TEMPLATE_PATH}`. "
                        "Vui lòng tạo thư mục `templates/` cạnh app.py và đặt file mẫu công ty vào đó "
                        "với tên `BaoCaoTuan_Template.docx`."
                    )
                else:
                    try:
                        docx_buffer = generate_weekly_report_from_template(TEMPLATE_PATH, report_ctx)
                        st.session_state["weekly_report_buffer"] = docx_buffer.getvalue()
                        st.session_state["weekly_report_filename"] = f"BaoCao_Tuan{report_week}_Thang{report_month}_{report_year}_MauCongTy.docx"
                        st.success("Đã tạo file theo mẫu công ty. Bấm nút bên dưới để tải về.")
                    except ValueError as e:
                        st.error(f"Không khớp được cấu trúc file mẫu: {e}")

        with exp_col2:
            st.caption("Bản tổng quát đơn giản do hệ thống tự dựng (không theo layout công ty) — dùng khi chưa có file mẫu.")
            if st.button("📋 Xuất bản tổng quát (.docx)", use_container_width=True):
                docx_buffer = generate_weekly_report_docx(report_ctx)
                st.session_state["weekly_report_buffer"] = docx_buffer.getvalue()
                st.session_state["weekly_report_filename"] = f"BaoCao_Tuan{report_week}_Thang{report_month}_{report_year}_TongQuat.docx"
                st.success("Đã tạo file báo cáo. Bấm nút bên dưới để tải về.")

        if st.session_state.get("weekly_report_buffer"):
            st.download_button(
                "⬇️ Tải file báo cáo Word",
                data=st.session_state["weekly_report_buffer"],
                file_name=st.session_state.get("weekly_report_filename", "BaoCao.docx"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

else:
    st.markdown(f'<div class="main-header">{page}</div>', unsafe_allow_html=True)
    st.info("🚧 Phân hệ này đang được thiết kế giao diện.")
